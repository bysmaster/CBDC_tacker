# -*- coding: utf-8 -*-
import os
import smtplib
from datetime import datetime
from email.header import Header
from email.mime.application import MIMEApplication
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from pathlib import Path
from typing import List, Dict, Tuple, Optional

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from src.utils import to_chinese_numeral

# Configuration for Email
EMAIL_HOST = "smtp.qq.com"
EMAIL_PORT = 465
EMAIL_USER = os.environ.get("EMAIL_USER")
EMAIL_PASS = os.environ.get("EMAIL_PASS")
EMAIL_TO = os.environ.get("EMAIL_TO")

ROOT_DIR = Path(__file__).resolve().parents[2] # src/pipeline/ -> src/ -> root
DATA_DIR = ROOT_DIR / "data"
TEMPLATE_PATH = ROOT_DIR / "memo" / "模板.docx"
OUTPUT_DOC_DIR = DATA_DIR / "reports"
OUTPUT_DOC_DIR.mkdir(exist_ok=True, parents=True)

def inspect_template_paragraphs(docx_path: str) -> List[str]:
    """Return a list of text from all paragraphs in the document."""
    try:
        doc = Document(docx_path)
        return [p.text for p in doc.paragraphs]
    except Exception as e:
        print(f"Error reading docx: {e}")
        return []

def validate_template_integrity(docx_path: str) -> Tuple[bool, str]:
    path = Path(docx_path)
    if not path.exists():
        return False, f"Template file not found: {docx_path}"

    paragraphs = inspect_template_paragraphs(str(path))
    
    required_markers = [
        "新疆维吾尔自治区分行",
        "编译："
    ]
    
    missing = []
    for marker in required_markers:
        if not any(marker in p_text for p_text in paragraphs):
            missing.append(marker)
            
    if missing:
        return False, f"Missing critical markers in template: {missing}"
        
    return True, "Template is valid."

def set_run_font(run, font_name, size_pt, bold=False):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    r = run._element
    rPr = r.get_or_add_rPr()
    fonts = rPr.get_or_add_rFonts()
    fonts.set(qn('w:eastAsia'), font_name)
    fonts.set(qn('w:ascii'), font_name)
    fonts.set(qn('w:hAnsi'), font_name)
    fonts.set(qn('w:cs'), font_name)

def generate_word_report(articles: List[Dict], date_str: str, filename_prefix: str = "CBDC_Report", filter_relevant: bool = True) -> Path:
    # Filter articles based on flag
    if filter_relevant:
        valid_articles = [a for a in articles if a.get('is_relevant') is True]
    else:
        valid_articles = articles
    
    # If no articles, we still might want to generate an empty report or handle it.
    # But for "All News", if 0 items, maybe no report.
    # User requested "Daily email regardless".
    # If valid_articles is empty, let's generate a "No Content" report to be safe, 
    # so the user sees an attachment confirming "Nothing today".
    
    is_valid, val_msg = validate_template_integrity(str(TEMPLATE_PATH))
    if not is_valid:
        raise ValueError(f"Template validation failed: {val_msg}")
    
    doc = Document(TEMPLATE_PATH)
    
    def format_paragraph_text(paragraph, text, font_name, size_pt, bold=False, align=None):
        paragraph.clear()
        run = paragraph.add_run(text)
        set_run_font(run, font_name, size_pt, bold)
        if align is not None:
            paragraph.alignment = align
        return run

    target_font_fs = "FangSong_GB2312"
    target_font_hei = "SimHei"
    
    # 1. Set Date
    for p in doc.paragraphs:
        if "2026年" in p.text and "月" in p.text:
            dt = datetime.now()
            parts = p.text.split("2026年")
            prefix = parts[0] if parts else "新疆维吾尔自治区分行········································"
            new_date_str = f"{dt.year}年{dt.month}月{dt.day}日"
            full_text = prefix + new_date_str
            format_paragraph_text(p, full_text, target_font_fs, 16)
            break
            
    # 2. Identify Sections
    date_idx = -1
    footer_idx = -1
    
    for i, p in enumerate(doc.paragraphs):
        if "新疆维吾尔自治区分行" in p.text:
            date_idx = i
        elif "编译：" in p.text:
            footer_idx = i
            break
            
    if date_idx != -1 and footer_idx != -1 and footer_idx > date_idx:
        # Clear content
        to_remove = []
        for i in range(date_idx + 1, footer_idx):
            to_remove.append(doc.paragraphs[i])
        for p in to_remove:
            p.clear()
            if p._element.getparent() is not None:
                p._element.getparent().remove(p._element)

        # Find footer again
        ref_p = None
        for p in doc.paragraphs:
            if "编译：" in p.text:
                ref_p = p
                break
        
        if ref_p:
            if not valid_articles:
                # Insert "No Content" message
                p_none = ref_p.insert_paragraph_before()
                format_paragraph_text(p_none, "今日无新增相关资讯。", target_font_fs, 16, align=WD_ALIGN_PARAGRAPH.CENTER)
                p_none.paragraph_format.space_after = Pt(24)
            else:
                for i, art in enumerate(valid_articles):
                    num_str = to_chinese_numeral(i + 1)
                    
                    # Title
                    p_title = ref_p.insert_paragraph_before()
                    # Use title_cn if available, else title
                    title_text = f"{num_str}、{art.get('title_cn', art.get('title', 'No Title'))}"
                    format_paragraph_text(p_title, title_text, target_font_hei, 16, bold=True)
                    p_title.paragraph_format.space_before = Pt(12)
                    p_title.paragraph_format.space_after = Pt(6)
                    
                    # Content
                    # For "All News", summary might be empty if not processed by AI?
                    # But processor runs AI for all. So summary should be there if relevant?
                    # If "All News" includes irrelevant, they might not have summary/title_cn populated nicely?
                    # Processor logic: AI runs for ALL items in CSV. 
                    # If irrelevant, AI still returns JSON with title_cn/summary/reasoning usually?
                    # Wait, processor.py: 
                    # if is_rel is True -> records result. 
                    # if irrelevant -> df.at...
                    # But 'result' object in loop has keys.
                    # We need to ensure 'summary' is present.
                    # If summary is empty, use abstract or first 200 chars of content.
                    
                    summary_text = art.get('summary', '')
                    if not summary_text:
                         summary_text = art.get('abstract', '') or art.get('content', '')[:200]
                    
                    p_content = ref_p.insert_paragraph_before()
                    format_paragraph_text(p_content, summary_text, target_font_fs, 16)
                    p_content.paragraph_format.first_line_indent = Pt(32)
                    p_content.paragraph_format.line_spacing = 1.5
                    p_content.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    
                    # Source
                    source_text = f"（{art.get('entity', 'Source')}：{art.get('url', '')}）"
                    p_source = ref_p.insert_paragraph_before()
                    format_paragraph_text(p_source, source_text, target_font_fs, 16, align=WD_ALIGN_PARAGRAPH.LEFT)
                    p_source.paragraph_format.left_indent = Pt(0)
                    p_source.paragraph_format.first_line_indent = Pt(0)
                    p_source.paragraph_format.space_after = Pt(12)

            if "编译：" in ref_p.text:
                format_paragraph_text(ref_p, ref_p.text, target_font_fs, 16)
            for p in doc.paragraphs:
                if "编审：" in p.text:
                    format_paragraph_text(p, p.text, target_font_fs, 16)
    
    filename = f"{filename_prefix}_{datetime.now().strftime('%Y%m%d')}.docx"
    out_path = OUTPUT_DOC_DIR / filename
    doc.save(out_path)
    return out_path

def send_email_with_attachment(filepaths: List[Path], csv_content_html: str = "", stats: Dict = None):
    if not all([EMAIL_USER, EMAIL_PASS, EMAIL_TO]):
        print("⚠️ Email credentials missing. Skipping email.")
        return

    msg = MIMEMultipart()
    msg['From'] = EMAIL_USER
    msg['To'] = EMAIL_TO
    msg['Subject'] = Header(f"数字货币国际资讯日报 - {datetime.now().strftime('%Y-%m-%d')}", 'utf-8')
    
    body_str = "<p>附件为今日数字货币国际资讯日报，请查收。</p>"
    
    if stats:
        body_str += f"""
        <h3>=== 今日监测情况 (Monitoring Status) ===</h3>
        <ul>
            <li>历史总条数 (Total History): {stats.get('total_all', 0)}</li>
            <li>今日新增条数 (New Today): {stats.get('total_new', 0)}</li>
            <li>CBDC相关条数 (CBDC Relevant): {stats.get('total_relevant', 0)}</li>
            <li>异常条数 (Errors): {stats.get('total_errors', 0)}</li>
        </ul>
        """
        
    if csv_content_html:
        body_str += "<h3>=== 资讯列表 (News List) ===</h3>"
        body_str += csv_content_html
        
    body = MIMEText(body_str, 'html', 'utf-8')
    msg.attach(body)
    
    # Attach all files
    for fp in filepaths:
        if fp and fp.exists():
            with open(fp, 'rb') as f:
                part = MIMEApplication(f.read(), Name=fp.name)
                part['Content-Disposition'] = f'attachment; filename="{fp.name}"'
                msg.attach(part)
        
    try:
        server = smtplib.SMTP_SSL(EMAIL_HOST, EMAIL_PORT)
        server.login(EMAIL_USER, EMAIL_PASS)
        server.sendmail(EMAIL_USER, EMAIL_TO, msg.as_string())
        server.quit()
        print(f"✅ Email sent successfully to {EMAIL_TO}")
    except Exception as e:
        print(f"❌ Email failed: {e}")
