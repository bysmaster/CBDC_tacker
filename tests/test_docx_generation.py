# -*- coding: utf-8 -*-
"""
Unit tests for document generation functionality.
Ensures python-docx is working correctly and compatible with project requirements.
"""

import unittest
import os
from pathlib import Path
from docx import Document

class TestDocxGeneration(unittest.TestCase):
    def setUp(self):
        self.test_file = Path("test_output.docx")
        if self.test_file.exists():
            os.remove(self.test_file)

    def tearDown(self):
        if self.test_file.exists():
            os.remove(self.test_file)

    def test_create_basic_document(self):
        """Test creating a basic Word document."""
        doc = Document()
        doc.add_heading('Test Document', 0)
        doc.add_paragraph('This is a test paragraph.')
        doc.save(self.test_file)
        
        self.assertTrue(self.test_file.exists(), "Document file was not created")
        
        # Verify content
        doc_read = Document(self.test_file)
        self.assertEqual(len(doc_read.paragraphs), 2)
        self.assertEqual(doc_read.paragraphs[0].text, 'Test Document')
        self.assertEqual(doc_read.paragraphs[1].text, 'This is a test paragraph.')

if __name__ == '__main__':
    unittest.main()
